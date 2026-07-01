#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将Markdown需求分析报告转换为格式化Word文档 (修复版)"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement

SRC = r"c:\Users\24924\Desktop\智能运输系统设计与集成综合实验\需求分析报告-城市交通流量预测与实时预警系统.md"
DST = r"c:\Users\24924\Desktop\智能运输系统设计与集成综合实验\第二组-城市交通流量预测与实时预警系统-需求分析报告.docx"

doc = Document()

# ── Page setup (A4) ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Define styles properly ──
# Normal
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(12)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style_normal.paragraph_format.line_spacing = 1.5

# Heading 1
try:
    h1 = doc.styles['Heading 1']
except KeyError:
    h1 = doc.styles.add_style('Heading 1', 1)
h1.font.name = '黑体'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)

# Heading 2
try:
    h2 = doc.styles['Heading 2']
except KeyError:
    h2 = doc.styles.add_style('Heading 2', 2)
h2.font.name = '黑体'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(4)

# Heading 3
try:
    h3 = doc.styles['Heading 3']
except KeyError:
    h3 = doc.styles.add_style('Heading 3', 3)
h3.font.name = '黑体'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(4)

# ── Helper functions ──
def add_cover():
    for _ in range(7):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智能运输系统设计与集成综合实验')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('城市交通流量预测与实时预警系统')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 需求分析报告 ——')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(5):
        doc.add_paragraph('')

    info_lines = [
        ('课程名称', '智能运输系统设计与集成综合实验'),
        ('选    题', '城市交通流量预测与实时预警系统'),
        ('小    组', '第二组'),
        ('完成日期', '2026年    月    日'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

def add_header_footer():
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run('智能运输系统设计与集成综合实验 — 需求分析报告')
        hr.font.size = Pt(9)
        hr.font.name = '宋体'
        hr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        hr.font.color.rgb = RGBColor(128, 128, 128)
        # bottom border for header
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # top border
        pPr = fp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '4')
        top.set(qn('w:color'), 'CCCCCC')
        pBdr.append(top)
        pPr.append(pBdr)

        run1 = fp.add_run('— ')
        run1.font.size = Pt(9)
        run1.font.name = '宋体'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # PAGE field
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run_page = fp.add_run()
        run_page._element.append(fld_begin)

        inst = OxmlElement('w:instrText')
        inst.set(qn('xml:space'), 'preserve')
        inst.text = ' PAGE '
        run_inst = fp.add_run()
        run_inst._element.append(inst)

        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run_sep = fp.add_run('1')
        run_sep.font.size = Pt(9)
        run_sep._element.insert(0, fld_sep)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run_end = fp.add_run()
        run_end._element.append(fld_end)

        run2 = fp.add_run(' —')
        run2.font.size = Pt(9)
        run2.font.name = '宋体'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_page_number_field(paragraph):
    """在段落中添加页码域（用于TOC后手动添加）"""
    pass  # not needed, footer handles it

def add_para(text, level=None):
    """添加段落，自动选择样式"""
    if level == 1:
        p = doc.add_paragraph(text, style='Heading 1')
    elif level == 2:
        p = doc.add_paragraph(text, style='Heading 2')
    elif level == 3:
        p = doc.add_paragraph(text, style='Heading 3')
    elif level == 4:
        p = doc.add_paragraph(text, style='Heading 3')
        if p.runs:
            p.runs[0].font.size = Pt(12)
    else:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code_block(code_text):
    """代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.5)

        pPr = pf.element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F0F0F0')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)

        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)

def add_table_from_md(rows_data):
    """从二维数组创建表格"""
    if not rows_data or len(rows_data) < 2:
        return

    ncols = max(len(r) for r in rows_data)

    # Normalize: ensure all rows have ncols columns
    for r in rows_data:
        while len(r) < ncols:
            r.append('')

    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to fit page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # Header row
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '2E75B6')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            # Zebra striping
            elif i % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F2F7FB')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)

    doc.add_paragraph('')

# ── Read Markdown ──
with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# ── Build document ──
add_cover()
add_header_footer()

# ── Parse and add content ──
i = 0
in_code_block = False
code_lines = []
in_table = False
table_rows = []
skip_toc = False

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip the YAML header block at top of file
    if i < 30:
        if line.startswith('# 城市') and '需求分析报告' in line:
            i += 1
            continue
        if line.startswith('> **课程**') or line.startswith('> **选题**') or line.startswith('> **阶段**'):
            i += 1
            continue
        if line == '---':
            i += 1
            continue

    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            if code_lines:
                add_code_block('\n'.join(code_lines))
                code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # Skip TOC listing (the - [摘要](#...) etc. section)
    if line.startswith('- [摘要]') or line.strip().startswith('- [摘要]'):
        skip_toc = True
        i += 1
        continue
    if skip_toc and (line.strip().startswith('1. [') or line.strip().startswith('2. [') or
                     line.strip().startswith('3. [') or line.strip().startswith('4. [') or
                     line.strip().startswith('5. [') or line.strip().startswith('6. [') or
                     line.strip().startswith('7. [') or line.strip().startswith('8. [') or
                     line.strip().startswith('9. [') or line.strip().startswith('- [') or
                     line.strip().startswith('  - [')):
        i += 1
        continue
    else:
        skip_toc = False

    # Headings
    if line.startswith('## '):
        add_para(line[3:], level=1)
        i += 1
        continue
    if line.startswith('### '):
        add_para(line[4:], level=2)
        i += 1
        continue
    if line.startswith('#### '):
        add_para(line[5:], level=3)
        i += 1
        continue

    # Table rows
    if '|' in line and line.strip().startswith('|'):
        if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
            i += 1
            continue

        if not in_table:
            in_table = True
            table_rows = []

        cells = [c.strip() for c in line.strip().split('|')]
        cells = [c for c in cells if c]
        table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_md(table_rows)
            table_rows = []
            in_table = False

    # Blockquote
    if line.strip().startswith('> '):
        text = line.strip()[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue

    # Horizontal rule
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '2E75B6')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # Regular text line
    if line.strip():
        p = doc.add_paragraph()
        # Handle **bold** spans
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            elif part:
                run = p.add_run(part)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        doc.add_paragraph('')

    i += 1

# Flush remaining table
if in_table and table_rows:
    add_table_from_md(table_rows)
if in_code_block and code_lines:
    add_code_block('\n'.join(code_lines))

# ── Remove the empty first "section" if it creates a blank page ──
# Save
try:
    doc.save(DST)
    print(f'Word document generated: {DST}')
    size = os.path.getsize(DST)
    print(f'File size: {size} bytes')
except Exception as e:
    print(f'Error saving: {e}')
    raise
